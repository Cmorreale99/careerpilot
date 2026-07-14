"""Tests for the deterministic profile markdown renderer (no DB needed)."""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from app.profile_builder import profile_docx, profile_markdown, summary_docx  # noqa: E402

PROFILE = {
    "subject": "Cam Morreale",
    "pipeline": {"documents": 2, "chunks": 10, "canonical_chunks": 9, "artifacts": 1},
    "resolved_references": [
        {"slug": "greenie", "canonical": "oneworld", "name": "Greenie"},
    ],
    "sections": [
        {"type": "professional_experience", "title": "Professional Experience",
         "artifacts": []},
        {"type": "project", "title": "Projects", "artifacts": [{
            "slug": "demo",
            "name": "Demo Project",
            "grounded_in": "Demo doc title.",
            "evidence_count": 2,
            "source_count": 1,
            "sources": [{
                "document": "2nd brain/demo.md",
                "source_file": "2nd brain/demo.pdf",
                "source_sha256": "abcdef0123456789",
                "lines": [
                    {"text": "Built the demo.", "page": 1, "method": "document", "matched": "x"},
                    {"text": "Shipped the demo.", "page": None, "method": "mention", "matched": "Demo"},
                ],
            }],
        }]},
    ],
}


class TestProfileMarkdown:
    def test_structure_and_content(self):
        md = profile_markdown(PROFILE)
        assert md.startswith("# Cam Morreale — Career Profile")
        assert "## Projects" in md
        assert "### Demo Project" in md
        assert "#### Source: `2nd brain/demo.pdf` (sha256 abcdef012345…)" in md
        assert "- Built the demo. (p. 1)" in md
        assert "- Shipped the demo.\n" in md  # no page tag when page is None

    def test_empty_sections_omitted(self):
        md = profile_markdown(PROFILE)
        assert "## Professional Experience" not in md

    def test_resolved_reference_note(self):
        md = profile_markdown(PROFILE)
        assert "“Greenie” (greenie) resolves to the oneworld artifact" in md

    def test_deterministic(self):
        assert profile_markdown(PROFILE) == profile_markdown(PROFILE)


SUMMARY_PROFILE = {
    "subject": "Cam Morreale",
    "kind": "summary",
    "models": ["qwen3.5:9b-q4_K_M"],
    "pipeline": {"documents": 2, "chunks": 10, "canonical_chunks": 9, "artifacts": 2},
    "sections": [
        {"type": "project", "title": "Projects", "artifacts": [
            {
                "slug": "demo", "name": "Demo Project", "grounded_in": "x",
                "summarized": True,
                "overview": "Built a demo system.",
                "contributions": [{"text": "Shipped it.", "chunk_ids": [4, 7]}],
                "capabilities": [{"text": "Python.", "chunk_ids": [4]}],
                "why_it_mattered": "Proved the idea.",
                "supporting_chunk_ids": [4, 7],
                "source_documents": ["2nd brain/demo.md"],
            },
            {
                "slug": "bare", "name": "Bare Artifact", "grounded_in": "y",
                "summarized": False,
                "overview": "", "contributions": [], "capabilities": [],
                "why_it_mattered": "", "supporting_chunk_ids": [],
                "source_documents": [],
            },
        ]},
        {"type": "education", "title": "Education", "artifacts": []},
    ],
}


class TestSummaryDocx:
    def test_structure_and_content(self):
        import io

        from docx import Document

        doc = Document(io.BytesIO(summary_docx(SUMMARY_PROFILE)))
        texts = [p.text for p in doc.paragraphs]
        styles = {p.text: p.style.name for p in doc.paragraphs}
        assert "Cam Morreale — Career Profile" in texts
        assert styles["Projects"] == "Heading 1"
        assert styles["Demo Project"] == "Heading 2"
        assert styles["Contributions"] == "Heading 3"
        assert styles["Shipped it. [4, 7]"] == "List Bullet"  # chunk IDs preserved
        assert "Why it mattered" in texts
        assert any(t.startswith("Sources: 2nd brain/demo.md") for t in texts)
        assert "Education" not in texts  # empty section omitted

    def test_unsummarized_artifact_points_to_corpus(self):
        import io

        from docx import Document

        doc = Document(io.BytesIO(summary_docx(SUMMARY_PROFILE)))
        texts = [p.text for p in doc.paragraphs]
        assert any("No generated summary" in t for t in texts)

    def test_intro_names_model_and_separate_corpus(self):
        import io

        from docx import Document

        doc = Document(io.BytesIO(summary_docx(SUMMARY_PROFILE)))
        intro = doc.paragraphs[1].text
        assert "qwen3.5:9b-q4_K_M" in intro
        assert "separate download" in intro


class TestProfileDocx:
    def test_structure_and_content(self):
        import io

        from docx import Document

        doc = Document(io.BytesIO(profile_docx(PROFILE)))
        texts = [p.text for p in doc.paragraphs]
        styles = {p.text: p.style.name for p in doc.paragraphs}
        assert "Cam Morreale — Career Profile" in texts
        assert styles["Projects"] == "Heading 1"
        assert styles["Demo Project"] == "Heading 2"
        assert any(t.startswith("Source: 2nd brain/demo.pdf (sha256 abcdef012345") for t in texts)
        assert styles["Built the demo. (p. 1)"] == "List Bullet"
        assert "Shipped the demo." in texts
        assert not any(t == "Professional Experience" for t in texts)  # empty section omitted
