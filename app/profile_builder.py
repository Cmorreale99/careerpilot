"""Deterministic career-profile assembly from the CareerPilot database.

Implementation step 9 of claude.md (assembly half): the comprehensive
profile is built entirely from Postgres — canonical artifacts and their
assigned canonical evidence chunks — with no LLM involved. Every
evidence line carries its source document, original file, page, and the
assignment method that placed it, so the rendered profile stays
traceable end to end.

Deterministic rules, documented here and encoded below:
  - sections in fixed order: professional experience, projects,
    hackathons, education
  - artifacts in registry order (artifacts.id) within a section;
    same_as duplicates are folded into their canonical artifact
  - evidence: canonical chunks only (duplicate_of IS NULL) of types
    sentence / list_item / table_row (headings structure the corpus and
    code blocks are implementation detail; both are excluded)
  - evidence grouped by source document (ordered by path), chunks in
    document order (chunk_index)
"""

from __future__ import annotations

import os
import sys
from pathlib import Path

import psycopg

REPO_ROOT = Path(__file__).resolve().parent.parent

SECTION_ORDER = [
    ("professional_experience", "Professional Experience"),
    ("project", "Projects"),
    ("hackathon", "Hackathons"),
    ("education", "Education"),
]

EVIDENCE_TYPES = ("sentence", "list_item", "table_row")


def database_url() -> str:
    url = os.environ.get("DATABASE_URL")
    if url:
        return url
    env_file = REPO_ROOT / ".env"
    if env_file.is_file():
        for line in env_file.read_text(encoding="utf-8").splitlines():
            key, _, value = line.strip().partition("=")
            if key == "DATABASE_URL" and value:
                return value
    print("error: DATABASE_URL not set and not found in .env", file=sys.stderr)
    raise SystemExit(1)


def build_profile() -> dict:
    """Assemble the full profile as a JSON-serializable dict."""
    with psycopg.connect(database_url()) as conn:
        doc_count, = conn.execute("SELECT count(*) FROM documents").fetchone()
        chunk_total, chunk_canonical = conn.execute(
            "SELECT count(*), count(*) FILTER (WHERE duplicate_of IS NULL) FROM chunks"
        ).fetchone()

        artifacts = conn.execute(
            """
            SELECT id, slug, name, artifact_type, grounded_in
            FROM artifacts WHERE same_as IS NULL ORDER BY id
            """
        ).fetchall()
        resolved = conn.execute(
            "SELECT slug, same_as, name FROM artifacts WHERE same_as IS NOT NULL ORDER BY slug"
        ).fetchall()

        evidence_rows = conn.execute(
            """
            SELECT ca.artifact_id, d.relative_path, d.source_file, d.source_sha256,
                   c.chunk_index, c.text_norm, c.source_page, ca.method, ca.matched_value
            FROM chunk_artifacts ca
            JOIN chunks c ON c.id = ca.chunk_id
            JOIN documents d ON d.id = c.document_id
            WHERE c.duplicate_of IS NULL AND c.chunk_type = ANY(%s)
            ORDER BY ca.artifact_id, d.relative_path, c.chunk_index
            """,
            (list(EVIDENCE_TYPES),),
        ).fetchall()

    by_artifact: dict[int, dict[str, dict]] = {}
    for (artifact_id, path, source_file, source_sha, index,
         text, page, method, matched) in evidence_rows:
        groups = by_artifact.setdefault(artifact_id, {})
        group = groups.setdefault(path, {
            "document": path,
            "source_file": source_file,
            "source_sha256": source_sha,
            "lines": [],
        })
        group["lines"].append({
            "text": text,
            "page": page,
            "method": method,
            "matched": matched,
        })

    sections = []
    for type_key, title in SECTION_ORDER:
        entries = []
        for artifact_id, slug, name, artifact_type, grounded_in in artifacts:
            if artifact_type != type_key:
                continue
            groups = list(by_artifact.get(artifact_id, {}).values())
            entries.append({
                "slug": slug,
                "name": name,
                "grounded_in": grounded_in,
                "evidence_count": sum(len(g["lines"]) for g in groups),
                "source_count": len(groups),
                "sources": groups,
            })
        sections.append({
            "type": type_key,
            "title": title,
            "artifacts": entries,
        })

    return {
        "subject": "Cam Morreale",
        "pipeline": {
            "documents": doc_count,
            "chunks": chunk_total,
            "canonical_chunks": chunk_canonical,
            "artifacts": len(artifacts),
        },
        "resolved_references": [
            {"slug": slug, "canonical": canonical, "name": name}
            for slug, canonical, name in resolved
        ],
        "sections": sections,
    }


def profile_markdown(profile: dict) -> str:
    """Render the profile dict as a downloadable Markdown document."""
    p = profile["pipeline"]
    out = [
        f"# {profile['subject']} — Career Profile",
        "",
        "Assembled deterministically from the CareerPilot evidence corpus: "
        f"{p['documents']} source documents → {p['chunks']} chunks → "
        f"{p['canonical_chunks']} after exact deduplication → "
        f"{p['artifacts']} artifacts. Every line below cites its source.",
        "",
    ]
    for section in profile["sections"]:
        if not section["artifacts"]:
            continue
        out += [f"## {section['title']}", ""]
        for artifact in section["artifacts"]:
            out += [
                f"### {artifact['name']}",
                "",
                f"*Grounding:* {artifact['grounded_in']}",
                "",
                f"*Evidence:* {artifact['evidence_count']} lines "
                f"from {artifact['source_count']} source documents.",
                "",
            ]
            for group in artifact["sources"]:
                origin = group["source_file"] or group["document"]
                sha = (group["source_sha256"] or "")[:12]
                out.append(f"#### Source: `{origin}`" + (f" (sha256 {sha}…)" if sha else ""))
                out.append("")
                for line in group["lines"]:
                    page = f" (p. {line['page']})" if line["page"] else ""
                    out.append(f"- {line['text']}{page}")
                out.append("")
    for ref in profile["resolved_references"]:
        out.append(
            f"> Reference note: “{ref['name']}” ({ref['slug']}) resolves to "
            f"the {ref['canonical']} artifact."
        )
    return "\n".join(out).rstrip() + "\n"


def profile_docx(profile: dict) -> bytes:
    """Render the profile dict as a downloadable .docx document."""
    import io

    from docx import Document

    doc = Document()
    p = profile["pipeline"]
    doc.add_heading(f"{profile['subject']} — Career Profile", level=0)
    doc.add_paragraph(
        "Assembled deterministically from the CareerPilot evidence corpus: "
        f"{p['documents']} source documents → {p['chunks']} chunks → "
        f"{p['canonical_chunks']} after exact deduplication → "
        f"{p['artifacts']} artifacts. Every line below cites its source."
    )

    for section in profile["sections"]:
        if not section["artifacts"]:
            continue
        doc.add_heading(section["title"], level=1)
        for artifact in section["artifacts"]:
            doc.add_heading(artifact["name"], level=2)
            grounding = doc.add_paragraph()
            grounding.add_run("Grounding: ").italic = True
            grounding.add_run(artifact["grounded_in"])
            stats = doc.add_paragraph()
            stats.add_run("Evidence: ").italic = True
            stats.add_run(
                f"{artifact['evidence_count']} lines from "
                f"{artifact['source_count']} source documents."
            )
            for group in artifact["sources"]:
                origin = group["source_file"] or group["document"]
                sha = (group["source_sha256"] or "")[:12]
                heading = f"Source: {origin}" + (f" (sha256 {sha}…)" if sha else "")
                doc.add_heading(heading, level=3)
                for line in group["lines"]:
                    page = f" (p. {line['page']})" if line["page"] else ""
                    doc.add_paragraph(f"{line['text']}{page}", style="List Bullet")

    for ref in profile["resolved_references"]:
        note = doc.add_paragraph()
        note.add_run(
            f"Reference note: “{ref['name']}” ({ref['slug']}) resolves to "
            f"the {ref['canonical']} artifact."
        ).italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


if __name__ == "__main__":
    import json
    print(json.dumps(build_profile(), indent=2, ensure_ascii=False)[:2000])
